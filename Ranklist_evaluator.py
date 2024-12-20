import pdfplumber
import csv
import tempfile
import os
from docx import Document
from docx.shared import Pt

def extract_rows_with_optional_filters(
    pdf_path, 
    main_term_index, 
    secondary_term_index, 
    rank_column_index, 
    college_column_index, 
    min_rank, 
    max_rank, 
    target_colleges, 
    main_term, 
    secondary_term=None
):
    temp_csv_path = tempfile.mktemp(suffix='.csv')
    
    with open(temp_csv_path, 'w', newline='', encoding='utf-8') as csvfile:
        csv_writer = csv.writer(csvfile)
        
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages):
                print(f"Processing page {page_num + 1} of {total_pages}")
                tables = page.extract_tables()
                
                for table_num, table in enumerate(tables):
                    if not table or len(table) < 2:
                        print(f"Skipping empty or malformed table on page {page_num + 1}, table {table_num + 1}.")
                        continue

                    try:
                        for row in table[1:]:  # Start directly from the first data row
                            # Normalize the row to handle multi-line text
                            normalized_row = [' '.join(str(cell).replace('\n', ' ').strip().split()) for cell in row]
                            full_row_text = ' '.join(normalized_row).lower()  # Combine for full-text search

                            # Check main_term and secondary_term
                            main_term_matches = main_term.lower() in (normalized_row[main_term_index].lower() if main_term_index < len(normalized_row) else "")
                            secondary_term_matches = True  # Default to True if secondary_term is None
                            if secondary_term and secondary_term_index is not None:
                                if secondary_term_index < len(normalized_row):
                                    secondary_value = normalized_row[secondary_term_index]
                                    secondary_term_matches = secondary_value and secondary_term.lower() in secondary_value.lower()
                                else:
                                    print(f"Skipping row: secondary_term_index out of bounds on page {page_num + 1}, table {table_num + 1}.")
                                    secondary_term_matches = False
                            
                            # Continue only if terms match
                            if not (main_term_matches and secondary_term_matches):
                                continue
                                
                            # Initialize flags for rank and college checks
                            rank_within_range = True
                            college_matches = True
                            
                            # Rank filtering if rank_column_index is specified
                            if rank_column_index is not None and rank_column_index < len(normalized_row):
                                try:
                                    rank = int(normalized_row[rank_column_index])  # Convert rank to integer
                                    rank_within_range = min_rank <= rank <= max_rank
                                except ValueError:
                                    print(f"Skipping row due to invalid rank value on page {page_num + 1}, table {table_num + 1}.")
                                    continue

                            # College filtering if college_column_index and target_colleges are specified
                            if college_column_index is not None and target_colleges and college_column_index < len(normalized_row):
                                college_name = normalized_row[college_column_index].lower()
                                college_matches = any(target_college.lower() in college_name for target_college in target_colleges)

                            # Write row if it meets term, rank, and college criteria
                            if rank_within_range and college_matches:
                                csv_writer.writerow(normalized_row)
                    except Exception as e:
                        print(f"Error processing table on page {page_num + 1}, table {table_num + 1}: {e}")

                page.flush_cache()

    return temp_csv_path

def save_to_docx_table(csv_path, output_path):
    doc = Document()
    doc.add_heading('Extracted Rows', 0)

    table = None
    row_count = 0

    with open(csv_path, 'r', newline='', encoding='utf-8') as csvfile:
        csv_reader = csv.reader(csvfile)

        for row in csv_reader:
            if table is None:
                table = doc.add_table(rows=1, cols=len(row))
                table.autofit = True
                hdr_cells = table.rows[0].cells
                for i, column_name in enumerate(row):
                    hdr_cells[i].text = str(column_name)
                    run = hdr_cells[i].paragraphs[0].runs[0]
                    run.font.size = Pt(10)
            else:
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row):
                    row_cells[i].text = str(cell_value) if cell_value is not None else ''
                    run = row_cells[i].paragraphs[0].runs[0]
                    run.font.size = Pt(10)

            row_count += 1
            if row_count % 100 == 0:
                print(f"Processed {row_count} rows")

    if table is None:
        print("No data to save.")
    else:
        doc.save(output_path)
        print(f"Extracted rows saved in table format to {output_path}")

    os.remove(csv_path)

# Example usage
pdf_path = 'sample.pdf'             # Replace with the path to your PDF
min_rank = 6000                         # Min rank for filtering
max_rank = 12000                        # Max rank for filtering
main_term = 'All India'                   # Use the term for filtering rows(ex here using All India)
secondary_term = '(NBEMS) PAEDIATRICS'    # Use any additional term for filtering rows(ex here using All India)
main_term_index = 10                    # In which number coulmn the main_term located(start column with 0)
secondary_term_index = 8                # Bypass secondary term filtering vy 'None'

temp_csv_path = extract_rows_with_optional_filters(
    pdf_path,
    main_term_index,
    secondary_term_index,
    rank_column_index=1,                #  Omit rank filtering by passing None
    college_column_index=7,             # Omit college filtering by passing None
    min_rank=min_rank,
    max_rank=max_rank,
    target_colleges=['West Bengal'],    # Target colleges or places to filter by
    main_term=main_term,
    secondary_term=secondary_term
)

output_docx_path = 'wb_paed.docx'
save_to_docx_table(temp_csv_path, output_docx_path)
